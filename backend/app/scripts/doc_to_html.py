#!/usr/bin/env python3
"""
Word 转 HTML 转换脚本
基于 Apache POI 的 Java 实现思路，使用 python-docx 实现
"""

import sys
import os
import argparse
from docx import Document


def docx_to_html(input_path, output_path):
    """
    将 Word 文档转换为 HTML
    """
    try:
        print(f"开始转换: {input_path} -> {output_path}")

        # 读取 Word 文档
        doc = Document(input_path)

        # 构建 HTML 内容
        html_content = []
        html_content.append("<!DOCTYPE html>")
        html_content.append("<html>")
        html_content.append("<head>")
        html_content.append('<meta charset="UTF-8">')
        html_content.append("<title>Converted Document</title>")
        html_content.append("<style>")
        html_content.append("body { font-family: Arial, sans-serif; margin: 20px; }")
        html_content.append("h1, h2, h3 { color: #333; }")
        html_content.append("p { line-height: 1.6; margin: 10px 0; }")
        html_content.append("table { border-collapse: collapse; width: 100%; margin: 10px 0; }")
        html_content.append("th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }")
        html_content.append("th { background-color: #f2f2f2; }")
        html_content.append("img { max-width: 100%; height: auto; }")
        html_content.append("</style>")
        html_content.append("</head>")
        html_content.append("<body>")

        # 处理文档内容
        for element in doc.element.body:
            if element.tag.endswith("p"):  # 段落
                paragraph = doc.paragraphs[doc.element.body.index(element)]
                if paragraph.text.strip():
                    html_content.append(f"<p>{escape_html(paragraph.text)}</p>")

            elif element.tag.endswith("tbl"):  # 表格
                table_index = [i for i, tbl in enumerate(doc.tables) if tbl._element == element][0]
                table = doc.tables[table_index]
                html_content.append("<table>")
                for row in table.rows:
                    html_content.append("<tr>")
                    for cell in row.cells:
                        html_content.append(f"<td>{escape_html(cell.text)}</td>")
                    html_content.append("</tr>")
                html_content.append("</table>")

        html_content.append("</body>")
        html_content.append("</html>")

        # 写入 HTML 文件
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(html_content))

        print(f"转换成功: {output_path}")
        return True

    except Exception as e:
        print(f"转换失败: {str(e)}")
        return False


def escape_html(text):
    """转义 HTML 特殊字符"""
    if not text:
        return ""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#39;")
    )


def main():
    parser = argparse.ArgumentParser(description="Word 转 HTML 转换工具")
    parser.add_argument("-i", "--input", required=True, help="输入文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出文件路径")

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误: 输入文件不存在 - {args.input}")
        sys.exit(1)

    success = docx_to_html(args.input, args.output)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
