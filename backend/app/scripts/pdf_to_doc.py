#!/usr/bin/env python3
import argparse
import sys
import os
import traceback
from pdfminer.high_level import extract_text
from docx import Document


def pdf_to_doc_pdfminer(pdf_path, doc_path):
    """使用 pdfminer 提取文本并转换为 Word 文档"""
    try:
        print(f"[INFO] 开始转换: {pdf_path} -> {doc_path}")

        # 使用 pdfminer 提取文本
        text = extract_text(pdf_path)

        if not text.strip():
            print("[WARNING] PDF 文件中没有提取到文本内容")
            # 创建一个空的 Word 文档
            doc = Document()
            doc.add_paragraph("此 PDF 文件没有可提取的文本内容")
            doc.save(doc_path)
            print(f"[INFO] 已创建空文档: {doc_path}")
            return True

        # 创建 Word 文档
        doc = Document()

        # 按段落分割文本
        paragraphs = text.split("\n\n")
        total_paragraphs = 0

        for para in paragraphs:
            para = para.strip()
            if para:
                # 处理长段落，避免 Word 文档格式问题
                if len(para) > 1000:
                    # 对长段落进行分割
                    chunks = [para[i : i + 800] for i in range(0, len(para), 800)]
                    for chunk in chunks:
                        if chunk.strip():
                            doc.add_paragraph(chunk.strip())
                            total_paragraphs += 1
                else:
                    doc.add_paragraph(para)
                    total_paragraphs += 1

        # 如果没有任何段落，添加提示
        if total_paragraphs == 0:
            doc.add_paragraph("未能从 PDF 中提取到有效的文本内容")

        # 保存文档
        doc.save(doc_path)
        print(f"[SUCCESS] 转换成功: {pdf_path} -> {doc_path}")
        print(f"[INFO] 提取了 {total_paragraphs} 个段落")
        return True

    except Exception as e:
        print(f"[ERROR] pdfminer 转换失败: {str(e)}")
        traceback.print_exc()
        return False


def pdf_to_doc_pdfplumber(pdf_path, doc_path):
    """备选方案：使用 pdfplumber 提取文本"""
    try:
        import pdfplumber

        print(f"[INFO] 使用 pdfplumber 转换: {pdf_path} -> {doc_path}")

        doc = Document()
        total_pages = 0
        total_paragraphs = 0

        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, 1):
                print(f"[INFO] 处理第 {page_num}/{total_pages} 页")

                # 提取文本
                text = page.extract_text()

                if text and text.strip():
                    paragraphs = text.split("\n\n")
                    for para in paragraphs:
                        para = para.strip()
                        if para:
                            doc.add_paragraph(para)
                            total_paragraphs += 1
                else:
                    # 如果没提取到文本，添加占位符
                    doc.add_paragraph(f"第 {page_num} 页 - 无文本内容")
                    total_paragraphs += 1

        if total_paragraphs == 0:
            doc.add_paragraph("未能从 PDF 中提取到文本内容")

        doc.save(doc_path)
        print(f"[SUCCESS] pdfplumber 转换成功: {pdf_path} -> {doc_path}")
        print(f"[INFO] 处理了 {total_pages} 页，提取了 {total_paragraphs} 个段落")
        return True

    except Exception as e:
        print(f"[ERROR] pdfplumber 转换失败: {str(e)}")
        return False


def pdf_to_doc_fitz(pdf_path, doc_path):
    """使用 PyMuPDF (fitz) 提取文本"""
    try:
        import fitz  # PyMuPDF

        print(f"[INFO] 使用 PyMuPDF 转换: {pdf_path} -> {doc_path}")

        doc = Document()
        total_pages = 0
        total_paragraphs = 0

        # 打开 PDF 文件
        pdf_document = fitz.open(pdf_path)
        total_pages = len(pdf_document)

        for page_num in range(total_pages):
            print(f"[INFO] 处理第 {page_num + 1}/{total_pages} 页")

            # 获取页面
            page = pdf_document[page_num]

            # 提取文本
            text = page.get_text()

            if text and text.strip():
                paragraphs = text.split("\n\n")
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        doc.add_paragraph(para)
                        total_paragraphs += 1
            else:
                doc.add_paragraph(f"第 {page_num + 1} 页 - 无文本内容")
                total_paragraphs += 1

        pdf_document.close()

        if total_paragraphs == 0:
            doc.add_paragraph("未能从 PDF 中提取到文本内容")

        doc.save(doc_path)
        print(f"[SUCCESS] PyMuPDF 转换成功: {pdf_path} -> {doc_path}")
        print(f"[INFO] 处理了 {total_pages} 页，提取了 {total_paragraphs} 个段落")
        return True

    except Exception as e:
        print(f"[ERROR] PyMuPDF 转换失败: {str(e)}")
        return False


def main():
    parser = argparse.ArgumentParser(description="PDF 转 Word 文档")
    parser.add_argument("-i", "--input", required=True, help="输入 PDF 文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出 Word 文件路径")

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"[ERROR] 输入文件不存在: {args.input}")
        sys.exit(1)

    if not args.input.lower().endswith(".pdf"):
        print(f"[ERROR] 输入文件不是 PDF 格式: {args.input}")
        sys.exit(1)

    print(f"[INFO] 输入文件: {args.input}")
    print(f"[INFO] 输出文件: {args.output}")
    print(f"[INFO] 文件大小: {os.path.getsize(args.input)} 字节")

    # 按优先级尝试不同的转换方法
    success = False

    # 方法1: 使用 pdfminer（最稳定）
    print("[INFO] 尝试方法1: pdfminer")
    success = pdf_to_doc_pdfminer(args.input, args.output)

    # 方法2: 使用 pdfplumber
    if not success:
        print("[INFO] 尝试方法2: pdfplumber")
        success = pdf_to_doc_pdfplumber(args.input, args.output)

    # 方法3: 使用 PyMuPDF
    if not success:
        print("[INFO] 尝试方法3: PyMuPDF")
        success = pdf_to_doc_fitz(args.input, args.output)

    if success:
        # 验证输出文件
        if os.path.exists(args.output) and os.path.getsize(args.output) > 0:
            print(f"[SUCCESS] 最终转换成功，输出文件大小: {os.path.getsize(args.output)} 字节")
            sys.exit(0)
        else:
            print("[ERROR] 输出文件创建失败或为空")
            sys.exit(1)
    else:
        print("[ERROR] 所有转换方法都失败了")
        print("[INFO] 请检查是否安装了必要的依赖:")
        print("  pip install pdfminer.six python-docx pdfplumber PyMuPDF")
        sys.exit(1)


if __name__ == "__main__":
    main()
